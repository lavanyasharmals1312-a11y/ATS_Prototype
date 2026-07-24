"""
resume_extractor.py
-------------------
Multi-strategy text extractor for PDF and DOCX resume files.

Strategy:
  PDF  → pdfplumber (text + all table cells per page)
  DOCX → python-docx body paragraphs + table cells + section headers/footers
         in document visual order via XML traversal
"""

from pathlib import Path
from loguru import logger

from app.services.protocols import TextExtractor  # noqa: F401


class ResumeExtractor:
    """Single responsibility: extract raw text from a resume file."""

    async def extract(self, path: str, mime_type: str) -> str:
        full_path = Path(path)
        if not full_path.is_absolute():
            from app.config import settings
            full_path = Path(settings.upload_dir) / path

        try:
            if mime_type == "application/pdf":
                return await self._extract_pdf(full_path)
            elif mime_type in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ):
                return await self._extract_docx(full_path)
            else:
                raise ValueError(f"Unsupported mime type: {mime_type}")
        except Exception as e:
            logger.warning(f"Text extraction failed for {path}: {e}")
            return ""

    # ── PDF ────────────────────────────────────────────────────────────────

    async def _extract_pdf(self, path: Path) -> str:
        """
        Extract text from PDF using pdfplumber with full table cell capture.

        For each page:
          1. Extract normal paragraph text (page.extract_text)
          2. Extract all table cell text to catch column/sidebar content
          3. Deduplicate lines that appear in both (tables repeat in text dump)
        """
        import pdfplumber

        page_chunks: list[str] = []

        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                lines_seen: set[str] = set()
                parts: list[str] = []

                # ── 1. Normal text stream ──────────────────────────────────
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text:
                    for line in text.splitlines():
                        stripped = line.strip()
                        if stripped:
                            lines_seen.add(stripped.lower())
                            parts.append(stripped)

                # ── 2. Table cell text (catches sidebars, skill grids) ─────
                try:
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            for cell in row:
                                if not cell:
                                    continue
                                cell_stripped = cell.strip()
                                if cell_stripped and cell_stripped.lower() not in lines_seen:
                                    lines_seen.add(cell_stripped.lower())
                                    parts.append(cell_stripped)
                except Exception as tbl_err:
                    logger.debug(f"Table extraction error on page {page_num}: {tbl_err}")

                if parts:
                    page_chunks.append("\n".join(parts))

        result = "\n\n".join(page_chunks)
        logger.debug(f"PDF extracted {len(result)} chars from {path.name}")
        return result

    # ── DOCX ───────────────────────────────────────────────────────────────

    async def _extract_docx(self, path: Path) -> str:
        """
        Extract text from DOCX in visual order, capturing:
          - Body paragraphs
          - Table cells (often missed entirely by naive iteration)
          - Section headers and footers
        """
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(path))
        parts: list[str] = []

        # ── 1. Body in visual order (paragraphs interleaved with tables) ──
        for block in doc.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

            if tag == "p":
                # Plain paragraph
                text = "".join(node.text or "" for node in block.iter(qn("w:t")))
                stripped = text.strip()
                if stripped:
                    parts.append(stripped)

            elif tag == "tbl":
                # Table — iterate rows → cells → paragraphs inside each cell
                for row in block.findall(f".//{qn('w:tr')}"):
                    cell_texts: list[str] = []
                    for cell in row.findall(f".//{qn('w:tc')}"):
                        cell_content = "".join(
                            node.text or "" for node in cell.iter(qn("w:t"))
                        ).strip()
                        if cell_content:
                            cell_texts.append(cell_content)
                    if cell_texts:
                        parts.append(" | ".join(cell_texts))

        # ── 2. Headers and footers (name/contact often live here) ──────────
        for section in doc.sections:
            for container in (section.header, section.footer):
                try:
                    for para in container.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
                    for table in container.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text = cell.text.strip()
                                if text:
                                    parts.append(text)
                except Exception:
                    pass  # Some docs have no header/footer — that's fine

        result = "\n".join(parts)
        logger.debug(f"DOCX extracted {len(result)} chars from {path.name}")
        return result